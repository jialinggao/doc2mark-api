import os
import io
import re
import tempfile
import subprocess
from typing import BinaryIO, List, Tuple, Optional, Dict, Any
from loguru import logger
from docx import Document
from docx.oxml.ns import qn
from markitdown import MarkItDown
from openai import OpenAI
import html2text
from app.config import settings
from app.models import ImageMode
from app.services.image_processor import image_processor


def clean_word_special_chars(markdown: str) -> str:
    """
    清理 Word 文档中的特殊字符

    Word 文档常包含 Unicode 专用区 (Private Use Area, 0xE000-0xF900)
    的字符，这些字符在 Markdown 中通常没有实际意义，需要清理。

    Args:
        markdown: 原始 Markdown 文本

    Returns:
        清理后的 Markdown 文本
    """
    cleaned = markdown

    # 清理 Unicode 专用区字符 (0xE000-0xF900)
    for code in range(0xE000, 0xF900):
        char = chr(code)
        if char in cleaned:
            cleaned = cleaned.replace(char, '')

    return cleaned


class DocxListParser:
    """
    DOCX 文档列表格式解析器

    主要功能：
    1. 解析 Word 文档中的列表编号定义（abstractNum、lvl 等）
    2. 根据定义生成对应的编号（中文数字、罗马数字等）
    3. 提取文档中的所有列表项
    """

    def __init__(self):
        """
        初始化列表解析器
        """
        self.num_id_to_abstract_id: Dict[int, int] = {}
        self.abstract_num_info: Dict[Tuple[int, int], Dict[str, Any]] = {}
        self.list_counters: Dict[int, int] = {}

    def parse_numbering(self, doc: Document) -> None:
        """
        解析文档中的编号定义

        Args:
            doc: Document 对象
        """
        self.num_id_to_abstract_id.clear()
        self.abstract_num_info.clear()
        self.list_counters.clear()

        numbering_part = doc.part.numbering_part
        if not numbering_part:
            return

        numbering_element = numbering_part._element

        # 解析 num 元素（关联 numId 和 abstractNumId）
        for num in numbering_element.findall(qn("w:num")):
            num_id_attr = num.get(qn("w:numId"))
            if num_id_attr is None:
                continue

            abstract_num_id_elem = num.find(qn("w:abstractNumId"))
            if abstract_num_id_elem is None:
                continue

            abstract_num_id_attr = abstract_num_id_elem.get(qn("w:val"))
            if abstract_num_id_attr is None:
                continue

            try:
                num_id = int(num_id_attr)
                abstract_num_id = int(abstract_num_id_attr)
                self.num_id_to_abstract_id[num_id] = abstract_num_id
                self.list_counters[num_id] = 0
            except (ValueError, TypeError):
                continue

        # 解析 abstractNum 元素（包含具体的列表格式定义）
        for abstract_num in numbering_element.findall(qn("w:abstractNum")):
            abstract_num_id_attr = abstract_num.get(qn("w:abstractNumId"))
            if abstract_num_id_attr is None:
                continue

            try:
                abstract_num_id = int(abstract_num_id_attr)
            except (ValueError, TypeError):
                continue

            # 解析每个层级的格式定义
            for lvl in abstract_num.findall(qn("w:lvl")):
                ilvl_attr = lvl.get(qn("w:ilvl"))
                if ilvl_attr is None:
                    continue

                try:
                    ilvl = int(ilvl_attr)
                except (ValueError, TypeError):
                    continue

                info = {}

                # 解析层级文本（如 "%1."）
                lvl_text_elem = lvl.find(qn("w:lvlText"))
                if lvl_text_elem is not None:
                    info['lvlText'] = lvl_text_elem.get(qn("w:val"), "")

                # 解析编号格式（如 "chineseNum1"、"roman" 等）
                num_fmt_elem = lvl.find(qn("w:numFmt"))
                if num_fmt_elem is not None:
                    info['numFmt'] = num_fmt_elem.get(qn("w:val"), "decimal")

                # 解析起始值
                start_elem = lvl.find(qn("w:start"))
                if start_elem is not None:
                    start_val = start_elem.get(qn("w:val"), "1")
                    try:
                        info['start'] = int(start_val)
                    except (ValueError, TypeError):
                        info['start'] = 1

                if info:
                    self.abstract_num_info[(abstract_num_id, ilvl)] = info

    def _convert_number(self, number: int, num_fmt: str) -> str:
        """
        根据格式转换数字

        Args:
            number: 数字
            num_fmt: 格式类型

        Returns:
            格式化后的字符串
        """
        num_fmt = num_fmt.lower()

        if num_fmt == "decimal":
            return str(number)
        elif num_fmt == "chinesenum1" or num_fmt == "chinesecounting":
            return self._to_chinese_number(number)
        elif num_fmt == "chinesenum2":
            return self._to_chinese_number_traditional(number)
        elif num_fmt == "roman":
            return self._to_roman(number).upper()
        elif num_fmt == "lowerroman":
            return self._to_roman(number).lower()
        elif num_fmt == "ordinal":
            return str(number) + "º"
        elif num_fmt == "cardtext":
            return self._to_chinese_number(number)
        elif num_fmt == "letter":
            return chr(ord('A') + number - 1)
        elif num_fmt == "lowerletter":
            return chr(ord('a') + number - 1)
        elif num_fmt == "ordinaltext":
            return self._to_chinese_number(number) + "号"
        else:
            return str(number)

    def _to_chinese_number(self, number: int) -> str:
        """
        阿拉伯数字转中文数字（简体中文，一、二、三...）

        Args:
            number: 阿拉伯数字

        Returns:
            中文数字字符串
        """
        if number <= 0:
            return str(number)

        chinese_nums = "零一二三四五六七八九十"

        if number <= 10:
            return chinese_nums[number]
        elif number < 100:
            tens = number // 10
            ones = number % 10
            if ones == 0:
                return chinese_nums[tens] + "十"
            else:
                if tens == 1:
                    return "十" + chinese_nums[ones]
                return chinese_nums[tens] + "十" + chinese_nums[ones]
        elif number < 1000:
            hundreds = number // 100
            remainder = number % 100
            result = chinese_nums[hundreds] + "百"
            if remainder > 0:
                if remainder < 10:
                    result += "零" + chinese_nums[remainder]
                else:
                    result += self._to_chinese_number(remainder)
            return result
        elif number < 10000:
            thousands = number // 1000
            remainder = number % 1000
            result = chinese_nums[thousands] + "千"
            if remainder > 0:
                if remainder < 100:
                    result += "零"
                result += self._to_chinese_number(remainder)
            return result
        else:
            return str(number)

    def _to_chinese_number_traditional(self, number: int) -> str:
        """
        阿拉伯数字转中文数字（繁体中文，壹、贰、叁...）

        Args:
            number: 阿拉伯数字

        Returns:
            中文数字字符串
        """
        chinese_nums = "零壹贰叁肆伍陆柒捌玖拾"
        if number <= 10:
            return chinese_nums[number]
        elif number < 20:
            return "拾" + chinese_nums[number - 10]
        else:
            return str(number)

    def _to_roman(self, number: int) -> str:
        """
        阿拉伯数字转罗马数字

        Args:
            number: 阿拉伯数字

        Returns:
            罗马数字字符串
        """
        roman_numerals = [
            (1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'),
            (100, 'C'), (90, 'XC'), (50, 'L'), (40, 'XL'),
            (10, 'X'), (9, 'IX'), (5, 'V'), (4, 'IV'), (1, 'I')
        ]
        result = []
        for value, numeral in roman_numerals:
            while number >= value:
                result.append(numeral)
                number -= value
            if number == 0:
                break
        return ''.join(result)

    def get_paragraph_list_number(self, paragraph) -> Optional[str]:
        """
        获取段落的列表编号

        Args:
            paragraph: 段落对象

        Returns:
            列表编号字符串（如 "一、"、"(1)" 等）
        """
        p_pr = paragraph._element.find(qn("w:pPr"))
        if p_pr is None:
            return None

        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            return None

        num_id_elem = num_pr.find(qn("w:numId"))
        if num_id_elem is None:
            return None

        num_id_val = num_id_elem.get(qn("w:val"))
        if num_id_val is None:
            return None

        try:
            num_id = int(num_id_val)
        except (ValueError, TypeError):
            return None

        ilvl_elem = num_pr.find(qn("w:ilvl"))
        if ilvl_elem is None:
            ilvl = 0
        else:
            ilvl_val = ilvl_elem.get(qn("w:val"))
            if ilvl_val is None:
                ilvl = 0
            else:
                try:
                    ilvl = int(ilvl_val)
                except (ValueError, TypeError):
                    return None

        if num_id not in self.num_id_to_abstract_id:
            return None

        abstract_num_id = self.num_id_to_abstract_id[num_id]
        key = (abstract_num_id, ilvl)

        if key not in self.abstract_num_info:
            return None

        info = self.abstract_num_info[key]
        lvl_text = info.get('lvlText', "")

        if not lvl_text:
            return None

        # 获取起始值和格式
        start = info.get('start', 1)
        num_fmt = info.get('numFmt', 'decimal')

        # 递增计数器并生成当前编号
        if num_id not in self.list_counters:
            self.list_counters[num_id] = 0
        self.list_counters[num_id] += 1
        current_num = start + self.list_counters[num_id] - 1

        # 转换数字格式
        converted_num = self._convert_number(current_num, num_fmt)

        # 替换占位符（如 "%1"）
        result = lvl_text.replace("%1", converted_num)

        # 处理多层级占位符（%2-%9）
        for i in range(2, 10):
            placeholder = f"%{i}"
            if placeholder in result:
                result = result.replace(placeholder, self._convert_number(i, num_fmt))

        return result

    def extract_all_list_items(self, doc: Document) -> List[Tuple[str, str]]:
        """
        提取文档中的所有列表项

        Args:
            doc: Document 对象

        Returns:
            列表项列表，每项为 (编号, 文本内容)
        """
        self.parse_numbering(doc)

        result = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                list_num = self.get_paragraph_list_number(paragraph)
                if list_num:
                    result.append((list_num, text))

        return result

    def extract_list_items_with_style(self, doc: Document) -> List[Tuple[str, str]]:
        """
        通过样式识别提取列表项（备用方案）

        Args:
            doc: Document 对象

        Returns:
            列表项列表，每项为 (编号, 文本内容)
        """
        result = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                style_name = paragraph.style.name if paragraph.style else ""
                if "List" in style_name or "列表" in style_name or "ListParagraph" in style_name:
                    p_pr = paragraph._element.find(qn("w:pPr"))
                    if p_pr is not None:
                        num_pr = p_pr.find(qn("w:numPr"))
                        if num_pr is not None:
                            num_id_elem = num_pr.find(qn("w:numId"))
                            if num_id_elem is not None:
                                num_id = int(num_id_elem.get(qn("w:val")))
                                if num_id not in self.list_counters:
                                    self.list_counters[num_id] = 0
                                self.list_counters[num_id] += 1

                                # 尝试从文本开头提取编号
                                pattern = r"^([第第]?[\u4e00-\u9fff\d]+[条章节款项]?)\s*"
                                match = re.match(pattern, text)
                                if match:
                                    result.append((match.group(1), text))
                                else:
                                    result.append((str(self.list_counters[num_id]), text))
        return result


class WordConverter:
    """
    Word 文档转换器

    主要功能：
    1. 支持 .doc (旧版 Word) 和 .docx (新版 Word) 格式
    2. 支持 HTML 包装的 .doc 格式（MIME + HTML）
    3. 使用 LibreOffice 转换旧版格式为新版
    4. 使用 MarkItDown 或 html2text 进行文档转换
    5. 识别和恢复自定义列表格式（中文数字、罗马数字等）
    6. 处理图片提取和转换
    7. 清理 Word 特殊字符
    """

    OLD_TO_NEW_FORMAT = {".doc": ".docx"}

    def __init__(self):
        """
        初始化 Word 转换器
        """
        self.md = self._initialize_markitdown()
        self.docx_list_parser = DocxListParser()
        self.html2text_converter = self._initialize_html2text()

    def _initialize_html2text(self):
        """
        初始化 html2text 转换器

        Returns:
            html2text.HTML2Text 对象
        """
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.ignore_emphasis = False
        h.body_width = 0  # 不自动换行
        h.unicode_snob = True  # 使用 Unicode
        return h

    def _is_html_wrapped_doc(self, file_stream: BinaryIO) -> Tuple[bool, Optional[str]]:
        """
        检测是否是 HTML 包装的 .doc 文档

        Args:
            file_stream: 文件流

        Returns:
            (是否是 HTML 包装, 提取的 HTML 内容)
        """
        file_stream.seek(0)
        # 读取前 10KB 来检测
        content = file_stream.read(10240).decode('utf-8', errors='ignore')
        file_stream.seek(0)

        # 检测是否是 MIME 包装的 HTML
        has_mime_version = 'Mime-Version:' in content
        has_multipart = 'Content-Type: Multipart/related' in content
        has_html_doctype = '<!DOCTYPE html' in content.lower()
        has_office_namespace = 'xmlns:w="urn:schemas-microsoft-com:office:word"' in content

        is_html_wrapped = (has_mime_version and has_multipart) or (has_html_doctype and has_office_namespace)

        if is_html_wrapped:
            # 读取完整内容
            file_stream.seek(0)
            full_content = file_stream.read().decode('utf-8', errors='ignore')
            file_stream.seek(0)

            # 提取 HTML 内容
            html_content = self._extract_html_from_mime(full_content)
            return True, html_content

        return False, None

    def _extract_html_from_mime(self, content: str) -> str:
        """
        从 MIME 格式中提取 HTML 内容

        Args:
            content: MIME 格式的内容

        Returns:
            提取的 HTML 内容
        """
        # 查找 HTML 部分开始
        html_start = content.find('<!DOCTYPE html')
        if html_start == -1:
            html_start = content.find('<html')

        if html_start != -1:
            # 先取从 HTML 开始的内容
            html_part = content[html_start:]
            
            # 查找 HTML 结束标签，在第一个 MIME 边界处截断
            html_end_tag = html_part.find('</html>')
            if html_end_tag != -1:
                # 包含 </html> 标签
                html_content = html_part[:html_end_tag + len('</html>')]
                return html_content
            
            # 如果没找到 </html>，查找第一个 MIME 边界
            mime_boundary = html_part.find('--NEXT.ITEM-BOUNDARY')
            if mime_boundary != -1:
                html_content = html_part[:mime_boundary]
                return html_content
            
            return html_part

        return content

    def _clean_html_content(self, html_content: str) -> str:
        """
        清理 HTML 内容，移除不需要的元素

        Args:
            html_content: 原始 HTML 内容

        Returns:
            清理后的 HTML 内容
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, 'html.parser')

        # 移除不需要的元素
        elements_to_remove = [
            '.share',
            '.sets',
            '.func',
            '.wechat-qrcode',
            '.article-qrcode',
            '#share-1',
            '#Canvas',
            '.num_fs',
            '.actfwzh',
            '.gzh',
            '.gudownload',
            '.bot-btns-box',
            '.print',
            '.xxgk-download-box',
            '.jcjy',
            '#actdy',
            '#zwxz',
            'script',
            'style'
        ]

        for selector in elements_to_remove:
            for elem in soup.select(selector):
                elem.decompose()

        # 移除 style 属性（html2text 不需要）
        for elem in soup.find_all(style=True):
            del elem['style']

        # 移除空的 div
        for elem in soup.find_all('div'):
            if not elem.get_text(strip=True) and not elem.find():
                elem.decompose()

        return str(soup)

    def _convert_html_to_markdown(self, html_content: str) -> str:
        """
        将 HTML 转换为 Markdown

        Args:
            html_content: HTML 内容

        Returns:
            Markdown 内容
        """
        # 清理 HTML
        try:
            cleaned_html = self._clean_html_content(html_content)
        except ImportError:
            logger.warning("[Word转换] BeautifulSoup 未安装，跳过 HTML 清理")
            cleaned_html = html_content

        # 转换为 Markdown
        markdown = self.html2text_converter.handle(cleaned_html)

        # 清理多余的空行
        markdown = re.sub(r'\n{3,}', '\n\n', markdown)

        return markdown

    def _initialize_markitdown(self) -> MarkItDown:
        """
        初始化 MarkItDown 实例

        Returns:
            MarkItDown 对象
        """
        if settings.ENABLE_LLM:
            client = OpenAI(
                base_url=settings.LLM_BASE_URL,
                api_key=settings.LLM_API_KEY or "sk-local-model"
            )
            return MarkItDown(
                llm_client=client,
                llm_model=settings.LLM_MODEL,
                llm_prompt=settings.LLM_PROMPT
            )
        else:
            return MarkItDown()

    def convert_old_office_format(self, file_stream: BinaryIO, filename: str, new_ext: str) -> BinaryIO:
        """
        使用 LibreOffice 转换旧版 Office 格式

        Args:
            file_stream: 文件流
            filename: 文件名
            new_ext: 目标扩展名

        Returns:
            转换后的文件流
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, filename)
            with open(input_path, "wb") as f:
                f.write(file_stream.read())

            cmd = [
                "libreoffice", "--headless", "--convert-to", new_ext.lstrip("."),
                "--outdir", tmpdir, input_path
            ]

            logger.info(f"[Word转换] 执行命令: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice转换失败: {result.stderr}")

            base_name = os.path.splitext(filename)[0]
            output_path = os.path.join(tmpdir, base_name + new_ext)

            if not os.path.exists(output_path):
                raise RuntimeError(f"转换后的文件不存在: {output_path}")

            with open(output_path, "rb") as f:
                return io.BytesIO(f.read())

    def extract_docx_list_items(self, file_stream: BinaryIO) -> List[Tuple[str, str]]:
        """
        从 DOCX 文件中提取列表项

        Args:
            file_stream: 文件流

        Returns:
            列表项列表，每项为 (编号, 文本内容)
        """
        docx_list_items = []
        try:
            file_stream.seek(0)
            doc = Document(file_stream)
            docx_list_items = self.docx_list_parser.extract_all_list_items(doc)

            # 如果第一种方法失败，尝试通过样式识别
            if len(docx_list_items) == 0:
                docx_list_items = self.docx_list_parser.extract_list_items_with_style(doc)

            file_stream.seek(0)
        except Exception as e:
            logger.warning(f"[Word转换] DOCX列表解析失败: {e}")
            docx_list_items = []

        return docx_list_items

    def restore_custom_list_format(self, markdown: str, list_items: List[Tuple[str, str]]) -> str:
        """
        恢复自定义列表格式

        将 MarkItDown 转换后的默认列表格式（如 "1. "、"- "）
        替换为原始 Word 文档中的自定义格式（如 "一、"、"(1)" 等）

        Args:
            markdown: 原始 Markdown
            list_items: 从 DOCX 提取的列表项

        Returns:
            恢复后的 Markdown
        """
        if not list_items:
            return markdown

        lines = markdown.split('\n')
        new_lines = []
        item_index = 0

        for line in lines:
            line_stripped = line.strip()

            if not line_stripped:
                new_lines.append(line)
                continue

            if item_index < len(list_items):
                lvl_text, original_content = list_items[item_index]
                content_stripped = original_content.strip()

                if not content_stripped:
                    new_lines.append(line)
                    continue

                # 移除 Markdown 中默认的数字前缀
                line_clean = re.sub(r'^\d+\.\s*', '', line_stripped)

                # 匹配内容（忽略标点符号等）
                content_clean = re.sub(r'[^\w\u4e00-\u9fff]', '', content_stripped)
                line_clean_for_match = re.sub(r'[^\w\u4e00-\u9fff]', '', line_clean)

                min_match_len = min(len(content_clean), 8)
                if min_match_len >= 4 and line_clean_for_match.startswith(content_clean[:min_match_len]):
                    # 内容匹配，使用自定义列表格式
                    new_lines.append(f"- {lvl_text} {line_clean}")
                    item_index += 1
                elif re.match(r'^\d+\.\s*', line_stripped):
                    # 这是一个编号列表，尝试匹配
                    if item_index < len(list_items):
                        lvl_text, _ = list_items[item_index]
                        new_lines.append(f"- {lvl_text} {line_clean}")
                        item_index += 1
                    else:
                        new_lines.append(line)
                else:
                    new_lines.append(line)
            else:
                new_lines.append(line)

        return '\n'.join(new_lines)

    def convert(
        self,
        file_stream: BinaryIO,
        filename: str,
        enable_ocr: bool = False,
        enable_llm: bool = False,
        image_mode: ImageMode = ImageMode.BASE64,
        image_quality: int = 100,
        max_image_size: int = -1
    ) -> dict:
        """
        转换 Word 文档为 Markdown

        Args:
            file_stream: 文件二进制流
            filename: 文件名
            enable_ocr: 是否启用 OCR
            enable_llm: 是否启用 LLM
            image_mode: 图片输出模式
            image_quality: 图片质量
            max_image_size: 图片最大尺寸

        Returns:
            转换结果字典
        """
        import time
        start_time = time.time()
        ext = os.path.splitext(filename.lower())[1]

        # 先检测是否是 HTML 包装的文档
        is_html_wrapped, html_content = self._is_html_wrapped_doc(file_stream)
        if is_html_wrapped and html_content:
            logger.info(f"[Word转换] 检测到 HTML 包装的 .doc 文档，使用 HTML 转 Markdown 流程")
            raw_markdown = self._convert_html_to_markdown(html_content)
            images = []
        else:
            # 旧版 Word 格式转换
            current_stream = file_stream
            current_filename = filename
            if ext in self.OLD_TO_NEW_FORMAT:
                new_ext = self.OLD_TO_NEW_FORMAT[ext]
                current_filename = filename[:-len(ext)] + new_ext
                logger.info(f"[Word转换] 检测到旧版Word格式 {ext}，自动转换为 {new_ext}")
                current_stream = self.convert_old_office_format(file_stream, filename, new_ext)

            # 提取列表信息（仅对 .docx）
            docx_list_items = []
            if current_filename.lower().endswith('.docx'):
                docx_list_items = self.extract_docx_list_items(current_stream)

            # 使用 MarkItDown 转换
            result = self.md.convert(current_stream, file_path=current_filename, keep_data_uris=True)
            raw_markdown = result.text_content

            # 恢复列表格式
            if docx_list_items:
                raw_markdown = self.restore_custom_list_format(raw_markdown, docx_list_items)

            # 处理图片
            markdown_text, images = image_processor.process_markdown_images(
                raw_markdown, image_mode, image_quality, max_image_size, enable_ocr, enable_llm
            )

        # 清理特殊字符
        raw_markdown = clean_word_special_chars(raw_markdown)

        # 处理图片（如果是 HTML 流程，在这里处理）
        if is_html_wrapped and html_content:
            markdown_text, images = image_processor.process_markdown_images(
                raw_markdown, image_mode, image_quality, max_image_size, enable_ocr, enable_llm
            )

        duration = time.time() - start_time

        return {
            "filename": filename,
            "markdown": markdown_text,
            "images": images,
            "duration": round(duration, 2)
        }


# Word 转换器单例实例
word_converter = WordConverter()
