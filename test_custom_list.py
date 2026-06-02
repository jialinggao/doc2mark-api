#!/usr/bin/env python3
"""测试自定义列表格式保留功能"""

from docx import Document
from app.services.docx_list_parser import docx_list_parser
import sys

def test_docx_list_parsing(docx_path):
    """测试docx列表解析"""
    doc = Document(docx_path)
    
    print("=== 文档段落列表 ===")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            lvl_text = docx_list_parser.get_list_number_text(paragraph)
            if lvl_text:
                actual_num = docx_list_parser.get_paragraph_list_number(paragraph, doc)
                is_custom = docx_list_parser.is_custom_list_format(paragraph)
                print(f"段落 {i}: {lvl_text} -> {actual_num} (自定义: {is_custom})")
                print(f"    内容: {paragraph.text[:50]}...")
    
    custom_items = docx_list_parser.extract_custom_list_items(doc)
    print(f"\n=== 提取到 {len(custom_items)} 个自定义列表项 ===")
    for num, content in custom_items:
        print(f"{num}: {content[:50]}...")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python test_custom_list.py <docx文件路径>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    docx_list_parser.parse_numbering(Document(docx_path))
    test_docx_list_parsing(docx_path)